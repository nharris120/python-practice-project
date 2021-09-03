# FIRST CV DOCUMENT #

from docx import Document
from docx.shared import Inches
import pyttsx3


# USING SPEAK FUNCTION #
def speak(text):
    pyttsx3.speak(text)


document = Document()

# PROFILE PICTURE #
document.add_picture('me.jpg', width=Inches(2.0), height=Inches(2.0))

# NAME, NUMBER, EMAIL DETAILS #
name = input('What is your name? ')
speak('Hello' + name + 'How are you today?')

speak('What is your phone number?')
phone_number = input('What is your phone number? ')

speak('What is your email?')
email = input('What is your email?')

document.add_paragraph(
    name + ' | ' + phone_number + ' | ' + email)

# ABOUT ME SECTION #
document.add_heading('About Me')
document.add_paragraph(input('Tell me about your self? '))

# WORK EXPERIENCE #
document.add_heading('Work Experience')
p = document.add_paragraph()

company = input(' Enter company')
from_date = input('From Date')
to_date = input('To Date')

p.add_run(company + ' ').bold = True
p.add_run(from_date + '-' + to_date + '\n').italic = True

experience_details = input(
    'Describe your experience at' + company)
p.add_run(experience_details)

# ADD MORE EXPERIENCES #
while True:
    has_more_experiences = input(
        'Do you have more work experience? Yes or No ')
    if has_more_experiences.lower() == 'yes':
        p = document.add_paragraph()

        company = input(' Enter company' + ' ')
        from_date = input('From Date' + ' ')
        to_date = input('To Date' + ' ')

        p.add_run(company + ' ').bold = True
        p.add_run(from_date + '-' + to_date + '\n').italic = True

        experience_details = input(
            'Describe your experience at ' + company + ' ')
        p.add_run(experience_details)
    else:
        break

# ADDING SKILLS #
document.add_heading('Skills' + ' ')
skills = input('Enter Skill' + ' ')
p = document.add_paragraph(skills)
p.style = 'List Bullet'

while True:
    has_more_skills = input('Do you have more skills to add? Yes or No' + ' ')
    if has_more_skills.lower() == 'yes':
        skills = input('Enter Skill' + ' ')
        p = document.add_paragraph(skills)
        p.style = 'List Bullet'
    else:
        break

# FOOTER #
section = document.sections[0]
footer = section.footer
p = footer.paragraphs[0]
p.text = "CV generated using practice code"

document.save('cv.docx')
